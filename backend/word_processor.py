import io
import re
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =========================================================
# BẢNG ÁNH XẠ CHUYỂN ĐỔI BẢNG MÃ (TCVN3 -> UNICODE)
# =========================================================
TCVN3_CHARS = "¹|¸|®|ã|ä|å|æ|ç|é|ê|í|ì|î|ï|ó|ò|ñ|ô|õ|ö|÷|ø|ù|ú|û|ü|ý|þ|ÿ|đ|®|S|M|H|F"
UNICODE_CHARS = "á|à|ả|ã|ạ|â|ấ|ầ|ẩ|ẫ|ậ|ă|ắ|ằ|ẳ|ẵ|ặ|é|è|ẻ|ẽ|ẹ|ê|ế|ề|ể|ễ|ệ|í|ì|ỉ|ĩ|ị|ó|ò|ỏ|õ|ọ|ô|ố|ồ|ổ|ỗ|ộ|ơ|ớ|ờ|ở|ỡ|ợ|ú|ù|ủ|ũ|ụ|ư|ứ|ừ|ử|ữ|ự|ý|ỳ|ỷ|ỹ|ỵ|đ|Đ|S|M|H|F"

tcvn3_map = {
    'a\xcc\x81': 'á', 'a\xcc\x80': 'à', 'a\xcc\x89': 'ả', 'a\xcc\x83': 'ã', 'a\xcc\xa3': 'ạ',
    'â\xcc\x81': 'ấ', 'â\xcc\x80': 'ầ', 'â\xcc\x89': 'ẩ', 'â\xcc\x83': 'ẫ', 'â\xcc\xa3': 'ậ',
    'ă\xcc\x81': 'ắ', 'ă\xcc\x80': 'ằ', 'ă\xcc\x89': 'ẳ', 'ă\xcc\x83': 'ẵ', 'ă\xcc\xa3': 'ặ',
    'e\xcc\x81': 'é', 'e\xcc\x80': 'è', 'e\xcc\x89': 'ẻ', 'e\xcc\x83': 'ẽ', 'e\xcc\xa3': 'ẹ',
    'ê\xcc\x81': 'ế', 'ê\xcc\x80': 'ề', 'ê\xcc\x89': 'ể', 'ê\xcc\x83': 'ễ', 'ê\xcc\xa3': 'ệ',
    'i\xcc\x81': 'í', 'i\xcc\x80': 'ì', 'i\xcc\x89': 'ỉ', 'i\xcc\x83': 'ĩ', 'i\xcc\xa3': 'ị',
    'o\xcc\x81': 'ó', 'o\xcc\x80': 'ò', 'o\xcc\x89': 'ỏ', 'o\xcc\x83': 'õ', 'o\xcc\xa3': 'ọ',
    'ô\xcc\x81': 'ố', 'ô\xcc\x80': 'ồ', 'ô\xcc\x89': 'ổ', 'ô\xcc\x83': 'ỗ', 'ô\xcc\xa3': 'ộ',
    'ơ\xcc\x81': 'ớ', 'ơ\xcc\x80': 'ờ', 'ơ\xcc\x89': 'ở', 'ơ\xcc\x83': 'ỡ', 'ơ\xcc\xa3': 'ợ',
    'u\xcc\x81': 'ú', 'u\xcc\x80': 'ù', 'u\xcc\x89': 'ủ', 'u\xcc\x83': 'ũ', 'u\xcc\xa3': 'ụ',
    'ư\xcc\x81': 'ứ', 'ư\xcc\x80': 'ừ', 'ư\xcc\x89': 'ử', 'ư\xcc\x83': 'ữ', 'ư\xcc\xa3': 'ự',
    'y\xcc\x81': 'ý', 'y\xcc\x80': 'ỳ', 'y\xcc\x89': 'ỷ', 'y\xcc\x83': 'ỹ', 'y\xcc\xa3': 'ỵ',
    '®': 'Đ', '®': 'đ',
    # Thêm ánh xạ thủ công một số chữ TCVN3 phổ biến lỗi (.VnTime)
    '¸': 'á', 'µ': 'à', '¶': 'ả', '·': 'ã', '¹': 'ạ',
    'Ê': 'ấ', 'Ç': 'ầ', 'È': 'ẩ', 'É': 'ẫ', 'Ë': 'ậ',
    '¾': 'ắ', '»': 'ằ', '¼': 'ẳ', '½': 'ẵ', 'Æ': 'ặ',
    'Ð': 'é', 'Ì': 'è', 'Î': 'ẻ', 'Ï': 'ẽ', 'Ñ': 'ẹ',
    'Õ': 'ế', 'Ò': 'ề', 'Ó': 'ể', 'Ô': 'ễ', 'Ö': 'ệ',
    'Ý': 'í', '×': 'ì', 'Ø': 'ỉ', 'Ü': 'ĩ', 'Þ': 'ị',
    'ã': 'ó', 'ß': 'ò', 'á': 'ỏ', 'â': 'õ', 'ä': 'ọ',
    'è': 'ố', 'å': 'ồ', 'æ': 'ổ', 'ç': 'ỗ', 'é': 'ộ',
    'í': 'ớ', 'ê': 'ờ', 'ë': 'ở', 'ì': 'ữ', 'î': 'ợ',
    'ó': 'ú', 'ò': 'ù', 'ô': 'ủ', 'õ': 'ũ', 'ö': 'ụ',
    'ø': 'ứ', '÷': 'ừ', 'ù': 'ử', 'ú': 'ữ', 'û': 'ự',
    'ý': 'ý', 'ú': 'ỳ', 'û': 'ỷ', 'ü': 'ỹ', 'þ': 'ỵ',
    '®': 'đ', '§': 'Đ'
}

def convert_tcvn3_to_unicode(text):
    """
    Hàm tự động chuyển đổi chuỗi chữ lỗi TCVN3 sang Unicode chuẩn dựng sẵn.
    """
    if not text:
        return text
    new_text = ""
    for char in text:
        new_text += tcvn3_map.get(char, char)
    return new_text

def has_chinese_characters(text):
    """
    Kiểm tra xem chuỗi văn bản có chứa ký tự tiếng Trung (Hanzi) hay không.
    Dải Unicode của tiếng Trung: \u4e00 đến \u9fff
    """
    if not text:
        return False
    return bool(re.search(r'[\u4e00-\u9fff]', text))

def fix_and_format_word(word_file, target_font="Times New Roman", font_size=13, 
                        line_spacing=1.15, space_after=6, auto_convert_tcvn3=False,
                        detect_chinese=False, chinese_font="SimSun"):
    """
    XỬ LÝ FILE WORD TOÀN DIỆN:
    - Sửa lỗi mã hóa (TCVN3 sang Unicode)
    - Nhận dạng tiếng Trung để đổi font riêng biệt (SimSun...)
    - Chuẩn hóa font chữ tiếng Việt/Anh (Times New Roman...) và căn lề.
    """
    doc = docx.Document(word_file)
    
    # 1. TỰ ĐỘNG CĂN CHỈNH LỀ ĐẸP ĐỂ IN (Chuẩn văn bản hành chính)
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.79)      # ~2.0 cm
        section.bottom_margin = Inches(0.79)   # ~2.0 cm
        section.left_margin = Inches(1.18)     # ~3.0 cm
        section.right_margin = Inches(0.79)    # ~2.0 cm

    # 2. XỬ LÝ TOÀN BỘ ĐOẠN VĂN BẢN
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.space_before = Pt(0)
        
        # Căn đều 2 bên với các đoạn văn dài
        if len(paragraph.text) > 80 and paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in paragraph.runs:
            # Sửa lỗi bảng mã TCVN3 nếu người dùng bật tùy chọn này
            if auto_convert_tcvn3:
                run.text = convert_tcvn3_to_unicode(run.text)

            # Lựa chọn Font chữ dựa trên nhận dạng ngôn ngữ
            current_font = target_font
            if detect_chinese and has_chinese_characters(run.text):
                current_font = chinese_font

            run.font.name = current_font
            run.font.size = Pt(font_size)
            
            # Ép Word nhận diện font đúng cách thông qua cấu trúc XML
            run._r.get_or_add_rPr().get_or_add_rFonts().set(docx.oxml.ns.qn('w:eastAsia'), current_font)
            run._r.get_or_add_rPr().get_or_add_rFonts().set(docx.oxml.ns.qn('w:hAnsi'), current_font)

    # 3. XỬ LÝ TOÀN BỘ BẢNG BIỂU (TABLES)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.space_after = Pt(3)
                    for run in paragraph.runs:
                        if auto_convert_tcvn3:
                            run.text = convert_tcvn3_to_unicode(run.text)

                        current_font = target_font
                        if detect_chinese and has_chinese_characters(run.text):
                            current_font = chinese_font

                        run.font.name = current_font
                        run.font.size = Pt(font_size - 1)
                        run._r.get_or_add_rPr().get_or_add_rFonts().set(docx.oxml.ns.qn('w:eastAsia'), current_font)
                        run._r.get_or_add_rPr().get_or_add_rFonts().set(docx.oxml.ns.qn('w:hAnsi'), current_font)

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream