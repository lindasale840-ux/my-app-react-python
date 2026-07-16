import docx
from docx.shared import Pt, RGBColor
import io
import re

def has_chinese_characters(text):
    if not text:
        return False
    return bool(re.search(r'[\u4e00-\u9fff]', text))

def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def mail_merge_docx(template_file, excel_file, target_font="Times New Roman", font_size=13,
                    detect_chinese=False, chinese_font="SimSun",
                    is_bold=False, is_italic=False, is_underline=False, text_color="#000000"):
    """
    HÀM TRỘN THƯ SIÊU BẢO VỆ CẤU TRÚC (TOKEN-LEVEL):
    - Khắc phục triệt để lỗi định dạng nhầm sang văn bản cố định đi kèm.
    - Chỉ tô đậm/tô màu KHỚP CHÍNH XÁC 100% với dữ liệu từ Excel đổ vào.
    """
    import pandas as pd
    
    font_mapping = {
        "SimSun": "宋体",
        "Microsoft YaHei": "微软雅黑",
        "KaiTi": "楷体",
        "SimHei": "黑体"
    }
    actual_chinese_font = font_mapping.get(chinese_font, "宋体") if detect_chinese else target_font

    df = pd.read_excel(excel_file)
    df.columns = [str(col).strip() for col in df.columns]
    output_files = []

    def replace_text_smart(paragraphs, context):
        for paragraph in paragraphs:
            combined_text = "".join([r.text for r in paragraph.runs])
            
            has_tag = False
            for key in context.keys():
                if f"{{{{ {key} }}}}" in combined_text or f"{{{{{key}}}}}" in combined_text:
                    has_tag = True
                    break
            
            if not has_tag:
                continue 

            # Thu thập danh sách các giá trị Excel thực tế cần điền cho dòng này
            filled_values = []
            for key, value in context.items():
                placeholder_space = f"{{{{ {key} }}}}"
                placeholder_no_space = f"{{{{{key}}}}}"
                
                if placeholder_space in combined_text or placeholder_no_space in combined_text:
                    # Đánh dấu thẻ đặc biệt trong chuỗi tạm thời để phân tách chính xác không bị lẫn
                    # Sử dụng ký tự đặc biệt 🌟 để bọc dữ liệu động lại
                    combined_text = combined_text.replace(placeholder_space, f"🌟{value}🌟")
                    combined_text = combined_text.replace(placeholder_no_space, f"🌟{value}🌟")
                    filled_values.append(value)

            # Xóa các runs cũ bị phân mảnh
            for r in paragraph.runs:
                r.text = ""
                
            # Dùng Regex cắt chuỗi theo ký tự bọc đặc biệt 🌟 để bóc tách text thường và text động
            # Sau đó tiếp tục bóc tách chữ tiếng Trung nếu có
            main_tokens = re.split(r'(🌟.*?🌟)', combined_text)
            
            for token in main_tokens:
                if not token:
                    continue
                
                is_dynamic = False
                # Nếu là khối text động được bọc bởi 🌟
                if token.startswith("🌟") and token.endswith("🌟"):
                    is_dynamic = True
                    token = token[1:-1] # Bỏ ký tự 🌟 đi để lấy text sạch
                    if not token:
                        continue
                
                # Cắt nhỏ tiếp khối này để xử lý rẽ nhánh font chữ tiếng Trung (nếu có)
                sub_tokens = re.split(r'([\u4e00-\u9fff]+)', token)
                
                for sub_token in sub_tokens:
                    if not sub_token:
                        continue
                        
                    new_run = paragraph.add_run(sub_token)
                    is_cn = detect_chinese and has_chinese_characters(sub_token)
                    
                    display_font_name = chinese_font if is_cn else target_font
                    current_font = actual_chinese_font if is_cn else target_font
                    
                    new_run.font.name = display_font_name
                    new_run.font.size = Pt(font_size)
                    
                    # ÁP ĐỊNH DẠNG NÂNG CAO: Chỉ áp dụng khi ĐÚNG là dữ liệu động từ Excel
                    if is_dynamic:
                        new_run.bold = is_bold
                        new_run.italic = is_italic
                        new_run.underline = is_underline
                        if text_color != "#000000":
                            new_run.font.color.rgb = hex_to_rgb(text_color)
                            
                    # Ép cấu trúc XML cho Word nhận diện chuẩn
                    rPr = new_run._r.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(docx.oxml.ns.qn('w:ascii'), display_font_name)
                    rFonts.set(docx.oxml.ns.qn('w:hAnsi'), display_font_name)
                    rFonts.set(docx.oxml.ns.qn('w:eastAsia'), current_font)
                    rFonts.set(docx.oxml.ns.qn('w:cs'), display_font_name)
                    
                    if is_cn:
                        rFonts.set(docx.oxml.ns.qn('w:hint'), 'eastAsia')
                    else:
                        rFonts.set(docx.oxml.ns.qn('w:hint'), 'default')

    for index, row in df.iterrows():
        template_file.seek(0)
        doc = docx.Document(template_file)
        context = {key: ("" if pd.isna(val) else str(val)) for key, val in row.to_dict().items()}
        
        replace_text_smart(doc.paragraphs, context)
        
        for table in doc.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    replace_text_smart(cell.paragraphs, context)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        first_col_val = str(list(context.values())[0]).replace("/", "_").replace("\\", "_")
        file_name = f"Tai_Lieu_{first_col_val}_{index + 1}.docx"
        
        output_files.append((file_name, file_stream))
        
    return output_files